"""
Export DOCX — Génère un document Word professionnel depuis un livrable Markdown

Format :
  - En-tête DataSphere Innovation (logo texte + titre)
  - Police : Calibri (Word standard)
  - Styles : Heading 1/2/3, Normal, Table, Code
  - Pied de page : page N/M + date + confidentiel
  - Métadonnées Word (auteur, société, titre)
"""

from __future__ import annotations
import io
import re
from datetime import datetime


def markdown_to_docx(
    title: str,
    content_markdown: str,
    author: str = "DataSphere Innovation",
    buyer_name: str | None = None,
    confidential: bool = True,
) -> bytes:
    """
    Convert a deliverable Markdown to a formatted Word (.docx) document.
    Returns raw bytes ready for HTTP response.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Document metadata ─────────────────────────────────────────────────────
    core = doc.core_properties
    core.author  = author
    core.company = "DataSphere Innovation"
    core.title   = title
    core.created = datetime.now()
    core.keywords = "DataSphere, Data Engineering, Livrable"

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)   # A4
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Cover page ────────────────────────────────────────────────────────────
    _add_cover(doc, title, buyer_name, author)

    # ── Table of contents placeholder ─────────────────────────────────────────
    doc.add_page_break()
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run("TABLE DES MATIÈRES")
    toc_run.bold = True
    toc_run.font.size = Pt(12)
    toc_para.paragraph_format.space_after = Pt(6)
    # Word TOC field
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), ' TOC \\o "1-3" \\h \\z \\u ')
    toc_para._p.addnext(fld)

    # ── Content ───────────────────────────────────────────────────────────────
    doc.add_page_break()
    _parse_markdown(doc, content_markdown)

    # ── Footer ────────────────────────────────────────────────────────────────
    _add_footer(doc, title, confidential)

    # ── Save to bytes ─────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_cover(doc, title: str, buyer_name: str | None, author: str):
    """Add a professional cover page."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Brand
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand.add_run("DataSphere Innovation")
    brand_run.font.size = Pt(14)
    brand_run.font.color.rgb = RGBColor(0xFA, 0xCC, 0x15)
    brand_run.bold = True
    brand.paragraph_format.space_before = Pt(80)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("IA Platform — Cabinet de conseil Data & IA")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Spacer
    for _ in range(5):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Main title
    main = doc.add_paragraph()
    main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_run = main.add_run(title)
    main_run.font.size = Pt(22)
    main_run.bold = True
    main_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    main.paragraph_format.space_after = Pt(20)

    # Client
    if buyer_name:
        client = doc.add_paragraph()
        client.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_run = client.add_run(f"À l'attention de : {buyer_name}")
        client_run.font.size = Pt(12)
        client_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        client.paragraph_format.space_after = Pt(40)

    # Metadata box
    for _ in range(4):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(
        f"Auteur : {author}  •  Date : {datetime.now().strftime('%d/%m/%Y')}  •  Version : 1.0"
    )
    info_run.font.size = Pt(9)
    info_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)


def _parse_markdown(doc, markdown: str):
    """Parse Markdown and add formatted paragraphs to the document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = markdown.split('\n')
    in_code_block = False
    in_table = False
    table_rows: list[list[str]] = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph(line, style='No Spacing')
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Table
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Skip separator rows (---|---|---)
            if not all(re.match(r'^[-:]+$', c) for c in cells if c):
                table_rows.append(cells)
            i += 1
            # Check if next line is still table
            if i >= len(lines) or not lines[i].strip().startswith('|'):
                _add_table(doc, table_rows)
                in_table = False
                table_rows = []
            continue

        if in_table and not line.strip().startswith('|'):
            _add_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Headings
        if line.startswith('#### '):
            _add_heading(doc, line[5:], 4)
        elif line.startswith('### '):
            _add_heading(doc, line[4:], 3)
        elif line.startswith('## '):
            _add_heading(doc, line[3:], 2)
        elif line.startswith('# '):
            _add_heading(doc, line[2:], 1)
        # Blockquote
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            _add_inline(p, line[2:])
            p.paragraph_format.left_indent = Pt(20)
        # Bullet
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_inline(p, line[2:])
        # Numbered list
        elif re.match(r'^\d+\. ', line):
            content = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            _add_inline(p, content)
        # Horizontal rule
        elif line.strip() in ('---', '***', '___'):
            p = doc.add_paragraph('_' * 60)
            p.runs[0].font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0) if p.runs else None
        # Code inline or empty
        elif line.strip():
            p = doc.add_paragraph()
            _add_inline(p, line)
        else:
            doc.add_paragraph()

        i += 1


def _add_heading(doc, text: str, level: int):
    """Add a styled heading."""
    from docx.shared import Pt, RGBColor

    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
    style = style_map.get(level, 'Heading 3')
    try:
        p = doc.add_paragraph(style=style)
    except Exception:
        p = doc.add_paragraph()
    _add_inline(p, text)

    # Custom color for headings
    colors = {1: RGBColor(0x0F, 0x17, 0x2A), 2: RGBColor(0x1E, 0x40, 0xAF),
              3: RGBColor(0x1E, 0x40, 0xAF), 4: RGBColor(0x47, 0x55, 0x69)}
    for run in p.runs:
        run.font.color.rgb = colors.get(level, RGBColor(0x0F, 0x17, 0x2A))


def _add_inline(para, text: str):
    """Parse inline markdown (bold, italic, code) and add runs to paragraph."""
    # Split on **bold**, *italic*, `code`
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            from docx.shared import Pt
            run.font.size = Pt(9)
        elif part:
            para.add_run(part)


def _add_table(doc, rows: list[list[str]]):
    """Add a formatted table to the document."""
    if not rows:
        return
    from docx.shared import Pt, RGBColor

    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= cols:
                break
            cell = table.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            _add_inline(para, cell_text)
            para.paragraph_format.space_after = Pt(0)
            # Header row styling
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Background color for header
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '1E40AF')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)

    doc.add_paragraph()  # spacer after table


def _add_footer(doc, title: str, confidential: bool):
    """Add page number footer."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.clear()

        # Left: title
        run_left = footer_para.add_run(f"{title[:40]}{'…' if len(title)>40 else ''}")
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        footer_para.add_run("  |  ")

        # Center: confidential
        if confidential:
            run_conf = footer_para.add_run("CONFIDENTIEL")
            run_conf.font.size = Pt(8)
            run_conf.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            run_conf.bold = True
            footer_para.add_run("  |  ")

        # Right: page number
        run_page = footer_para.add_run("Page ")
        run_page.font.size = Pt(8)
        run_page.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), ' PAGE ')
        run_fld = OxmlElement('w:r')
        fld.append(run_fld)
        footer_para._p.append(fld)
