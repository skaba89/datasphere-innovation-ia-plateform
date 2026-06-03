"""
CV Consultant Generator — produces a professional DOCX CV for a consultant profile.

Used for appels d'offres: each AO typically requires one CV per consultant proposed.
The generated CV follows the standard French public procurement format.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any


def generate_cv_docx(consultant: dict[str, Any], tender_context: dict[str, Any] | None = None) -> bytes:
    """
    Generate a professional consultant CV as DOCX bytes.

    consultant dict keys:
        name, title, summary, experience_years, daily_rate,
        skills (list of str), languages (list of str),
        experiences (list of {company, role, period, description}),
        education (list of {degree, school, year}),
        certifications (list of str)

    tender_context (optional):
        tender_title, buyer_name, reference
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ── Styles ────────────────────────────────────────────────────────────────
    GOLD = RGBColor(0x1E, 0x3A, 0x5F)      # Dark blue (professional)
    ACCENT = RGBColor(0xFA, 0xCC, 0x15)    # Gold accent
    DARK = RGBColor(0x1E, 0x29, 0x3B)
    GRAY = RGBColor(0x64, 0x74, 0x8B)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = DARK

    def add_heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper() if level == 1 else text)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11 if level == 1 else 10)
        run.font.color.rgb = GOLD

        # Add bottom border to H1
        if level == 1:
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1E3A5F")
            pBdr.append(bottom)
            pPr.append(pBdr)
        return p

    def add_para(text: str, bold: bool = False, italic: bool = False, size: int = 10, color: RGBColor = DARK):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = color
        return p

    def add_bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = DARK

    def add_kv(key: str, value: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        k_run = p.add_run(f"{key} : ")
        k_run.bold = True
        k_run.font.name = "Calibri"
        k_run.font.size = Pt(10)
        k_run.font.color.rgb = GRAY
        v_run = p.add_run(value)
        v_run.font.name = "Calibri"
        v_run.font.size = Pt(10)
        v_run.font.color.rgb = DARK

    # ── Header ─────────────────────────────────────────────────────────────────
    # Name
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(4)
    run_name = p_name.add_run(consultant.get("name", "Prénom NOM"))
    run_name.bold = True
    run_name.font.name = "Calibri"
    run_name.font.size = Pt(18)
    run_name.font.color.rgb = GOLD

    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run(consultant.get("title", "Consultant"))
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = GRAY
    run_title.italic = True

    # Meta info
    meta_parts = []
    if consultant.get("experience_years"):
        meta_parts.append(f"{consultant['experience_years']} ans d'expérience")
    if consultant.get("daily_rate"):
        meta_parts.append(f"TJM : {consultant['daily_rate']}")
    if meta_parts:
        add_para(" · ".join(meta_parts), size=9, color=GRAY)

    # Tender context banner
    if tender_context:
        p_banner = doc.add_paragraph()
        p_banner.paragraph_format.space_before = Pt(8)
        p_banner.paragraph_format.space_after = Pt(8)
        r = p_banner.add_run(
            f"CV établi pour : {tender_context.get('tender_title', '')} — {tender_context.get('buyer_name', '')} "
            f"({tender_context.get('reference', '')})"
        )
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
        r.italic = True

    doc.add_paragraph()

    # ── Summary ────────────────────────────────────────────────────────────────
    if consultant.get("summary"):
        add_heading("Résumé du profil")
        add_para(consultant["summary"])

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = consultant.get("skills", [])
    if skills:
        add_heading("Compétences techniques")
        # Group skills in lines of 4
        for i in range(0, len(skills), 5):
            chunk = skills[i:i + 5]
            add_para(" · ".join(chunk), color=DARK)

    # ── Languages ─────────────────────────────────────────────────────────────
    langs = consultant.get("languages", [])
    if langs:
        add_heading("Langues")
        add_para(" · ".join(langs))

    # ── Experiences ───────────────────────────────────────────────────────────
    experiences = consultant.get("experiences", [])
    if experiences:
        add_heading("Expériences professionnelles")
        for exp in experiences:
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_before = Pt(6)
            p_exp.paragraph_format.space_after = Pt(2)

            r_role = p_exp.add_run(f"{exp.get('role', 'Rôle')} — {exp.get('company', 'Client')}")
            r_role.bold = True
            r_role.font.name = "Calibri"
            r_role.font.size = Pt(10)
            r_role.font.color.rgb = DARK

            r_period = p_exp.add_run(f"  ({exp.get('period', '')})")
            r_period.font.name = "Calibri"
            r_period.font.size = Pt(9)
            r_period.font.color.rgb = GRAY
            r_period.italic = True

            if exp.get("description"):
                add_para(exp["description"], color=DARK)

            achievements = exp.get("achievements", [])
            for ach in achievements:
                add_bullet(ach)

    # ── Education ─────────────────────────────────────────────────────────────
    education = consultant.get("education", [])
    if education:
        add_heading("Formation")
        for edu in education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(2)
            p_edu.paragraph_format.space_after = Pt(2)
            r = p_edu.add_run(f"{edu.get('degree', '')} — {edu.get('school', '')}")
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            r.font.color.rgb = DARK
            if edu.get("year"):
                r2 = p_edu.add_run(f"  ({edu['year']})")
                r2.font.name = "Calibri"
                r2.font.size = Pt(9)
                r2.font.color.rgb = GRAY

    # ── Certifications ────────────────────────────────────────────────────────
    certifications = consultant.get("certifications", [])
    if certifications:
        add_heading("Certifications")
        for cert in certifications:
            add_bullet(cert)

    # ── Footer ─────────────────────────────────────────────────────────────────
    p_footer = doc.add_paragraph()
    p_footer.paragraph_format.space_before = Pt(20)
    r_footer = p_footer.add_run(
        f"Document généré par DataSphere Innovation IA Platform — {datetime.now(timezone.utc).strftime('%d/%m/%Y')}"
    )
    r_footer.font.name = "Calibri"
    r_footer.font.size = Pt(8)
    r_footer.font.color.rgb = GRAY
    r_footer.italic = True

    # ── Serialize ─────────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
